import pytest
from docx import Document
from anonymizer.anonymizer import Anonymizer
from anonymizer.doc_handler import DocProcessor

def test_full_anonymization():
    anonymizer = Anonymizer(config_path="config.json")
    processor = DocProcessor(anonymizer)
    text = "Ali VELİ, Hacı Bayram Veli Mah. Şehit Teğmen Kalmaz Cad. No:2"
    doc = Document()
    doc.add_paragraph(text)
    processor.anonymize_document(doc)
    assert "[AD_SOYAD_1]" in doc.paragraphs[0].text
    assert "[ADRES_1]" in doc.paragraphs[0].text
