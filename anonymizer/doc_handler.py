from docx import Document
from anonymizer.anonymizer import Anonymizer
from docx.oxml.shared import qn

class DocProcessor:
    def __init__(self, anonymizer: Anonymizer):
        self.anonymizer = anonymizer

    def anonymize_paragraph(self, paragraph):
        full_text = paragraph.text
        if not full_text.strip():
            return

        anonymized_text = self.anonymizer.anonymize_text(full_text)
        if anonymized_text == full_text:
            return

       
        #linkin 'tıklanabilir' özelliğini kaldırır ama metni korur.(eposta'lar için önemli)
        hyperlinks = paragraph._p.xpath('.//w:hyperlink')
        for hl in hyperlinks:
            for r in hl.xpath('.//w:r'):
                paragraph._p.append(r) # Metin parçalarını paragrafa taşı
            hl.getparent().remove(hl) # Link kabuğunu sil

        # Mevcut ilk run'ın stil özelliklerini yedekle
        saved_style = None
        if paragraph.runs:
            first_run = paragraph.runs[0]
            saved_style = {
                'bold': first_run.bold,
                'italic': first_run.italic,
                'underline': first_run.underline,
                'font_name': first_run.font.name,
                'font_size': first_run.font.size,
                'color': first_run.font.color.rgb if first_run.font.color else None
            }

       
        for run in paragraph.runs:
            run.text = ""

        # 4. Anonim metni uygula ve stili geri yükle
        new_run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
        new_run.text = anonymized_text
        
        if saved_style:
            new_run.bold = saved_style['bold']
            new_run.italic = saved_style['italic']
            new_run.underline = saved_style['underline']
            if saved_style['font_name']:
                new_run.font.name = saved_style['font_name']
            if saved_style['font_size']:
                new_run.font.size = saved_style['font_size']
            if saved_style['color']:
                new_run.font.color.rgb = saved_style['color']

    def anonymize_document(self, doc: Document) -> Document:
        for p in doc.paragraphs:
            self.anonymize_paragraph(p)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        self.anonymize_paragraph(p)
        return doc